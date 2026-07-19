"""
content_extraction.py — Advanced Document Content Extractor

Extracts the full readable content of any supported document and produces
a rich set of granular, hashable evidence blocks (paragraphs, sentences,
headings, tables, metadata fields, etc.) suitable for use in a Group Hash
authenticity scoring system.

Supported document types:
  • PDF  (native text via pypdf + layout via pdfminer, fallback to OCR)
  • DOCX (python-docx)
  • XLSX / XLS  (openpyxl / xlrd)
  • TXT / CSV / MD  (raw text)
  • Images  (OCR via pytesseract → pillow)

Produces three extractors that are each registered in EXTRACTOR_REGISTRY:
  1.  FullTextExtractor   — the complete verbatim document text, section-by-section
  2.  ContentBlockHasher  — individual SHA-256 hashes for every paragraph / cell /
                            sentence (used for granular Group Hash comparison)
  3.  LinguisticProfiler  — vocabulary richness, sentence statistics, readability
                            scores and per-sentence hashes (linguistic fingerprint)
"""

from __future__ import annotations

import hashlib
import io
import re
import unicodedata
from typing import Any, ClassVar, Dict, List, Optional, Tuple

from forensic_engine.base import BaseExtractor, compute_hashes
from forensic_engine.context import ExtractionContext
from forensic_engine.constants import (
    CATEGORY_DOCUMENT, RELIABILITY_HIGH, RELIABILITY_MEDIUM
)

# ─── optional heavy deps ────────────────────────────────────────────────────
try:
    import docx as _docx          # python-docx
    DOCX_OK = True
except ImportError:
    _docx = None
    DOCX_OK = False

try:
    import openpyxl as _openpyxl   # xlsx
    OPENPYXL_OK = True
except ImportError:
    _openpyxl = None
    OPENPYXL_OK = False

try:
    import xlrd as _xlrd           # old xls
    XLRD_OK = True
except ImportError:
    _xlrd = None
    XLRD_OK = False

try:
    from pdfminer.high_level import extract_text as _pdfminer_text   # noqa
    PDFMINER_TEXT_OK = True
except ImportError:
    _pdfminer_text = None
    PDFMINER_TEXT_OK = False

try:
    import pypdf as _pypdf
    PYPDF_OK = True
except ImportError:
    _pypdf = None
    PYPDF_OK = False

try:
    import pytesseract as _tess
    from PIL import Image as _PILImage
    TESSERACT_OK = True
except ImportError:
    _tess = None
    TESSERACT_OK = False


# ════════════════════════════════════════════════════════════════════════════
#  Internal helpers
# ════════════════════════════════════════════════════════════════════════════

def _sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8", errors="replace")).hexdigest()


def _normalise(text: str) -> str:
    """NFKC-normalise + collapse whitespace."""
    text = unicodedata.normalize("NFKC", text)
    return re.sub(r"\s+", " ", text).strip()


def _split_sentences(text: str) -> List[str]:
    """Naive but reliable sentence splitter (no NLTK dependency)."""
    parts = re.split(r"(?<=[.!?])\s+", text)
    return [p.strip() for p in parts if p.strip()]


def _split_paragraphs(text: str) -> List[str]:
    """Split on blank lines."""
    blocks = re.split(r"\n{2,}", text)
    return [b.strip() for b in blocks if b.strip()]


def _readability_score(text: str) -> Dict[str, Any]:
    """Flesch Reading Ease (no external lib)."""
    sentences = _split_sentences(text)
    words     = text.split()
    syllables  = sum(_count_syllables(w) for w in words)
    if not sentences or not words:
        return {"score": None, "grade": "N/A"}
    asl = len(words) / len(sentences)       # avg sentence length
    asw = syllables / len(words)            # avg syllables per word
    fre = 206.835 - 1.015 * asl - 84.6 * asw
    fre = max(0.0, min(100.0, round(fre, 2)))
    if fre >= 70:
        grade = "Easy"
    elif fre >= 50:
        grade = "Standard"
    elif fre >= 30:
        grade = "Difficult"
    else:
        grade = "Very Difficult"
    return {"flesch_reading_ease": fre, "grade": grade,
            "avg_sentence_length_words": round(asl, 2),
            "avg_syllables_per_word":    round(asw, 2)}


def _count_syllables(word: str) -> int:
    """Very rough syllable count for Flesch calculation."""
    word = word.lower().strip(".,;:!?\"'()-")
    if not word:
        return 0
    vowels = "aeiouy"
    count  = 0
    prev_vowel = False
    for ch in word:
        is_v = ch in vowels
        if is_v and not prev_vowel:
            count += 1
        prev_vowel = is_v
    if word.endswith("e"):
        count -= 1
    return max(1, count)


# ════════════════════════════════════════════════════════════════════════════
#  Per-format content readers
# ════════════════════════════════════════════════════════════════════════════

class _ContentReader:
    """
    Reads full document content into a normalised structure:
      {
        "full_text": str,           # the entire verbatim text
        "sections":  List[dict],    # [{title, text, level}, ...]
        "tables":    List[List[str]], # for spreadsheets / PDF tables
        "source":    str,           # how text was extracted
        "page_count": int | None,
      }
    """

    @classmethod
    def read(cls, context: ExtractionContext) -> Dict[str, Any]:
        ft = context.file_type
        # detect DOCX / XLSX from extension since mime can be generic
        ext = context.file_path.rsplit(".", 1)[-1].lower() if "." in context.file_path else ""

        if ext in ("docx", "doc"):
            return cls._read_docx(context)
        if ext in ("xlsx", "xls", "csv"):
            return cls._read_spreadsheet(context, ext)
        if ext in ("txt", "md", "rst", "html", "htm", "json", "xml", "log", "csv"):
            return cls._read_text(context)
        if ft == "pdf":
            return cls._read_pdf(context)
        if ft == "image":
            return cls._read_image(context)
        # fallback — try raw decode
        return cls._read_text(context)

    # ── PDF ─────────────────────────────────────────────────────────────────
    @staticmethod
    def _read_pdf(context: ExtractionContext) -> Dict[str, Any]:
        sections: List[Dict] = []
        full_text = ""
        source    = "unavailable"
        pages     = 0

        # Strategy 1: pypdf page-by-page
        if PYPDF_OK:
            try:
                reader = _pypdf.PdfReader(io.BytesIO(context.raw_data))
                pages  = len(reader.pages)
                parts  = []
                for i, page in enumerate(reader.pages):
                    txt = page.extract_text() or ""
                    txt = _normalise(txt)
                    if txt:
                        parts.append(txt)
                        sections.append({"title": f"Page {i + 1}", "text": txt, "level": 0})
                full_text = "\n\n".join(parts)
                source    = "pypdf"
            except Exception:
                pass

        # Strategy 2: pdfminer (richer layout, better for columns)
        if (not full_text or len(full_text) < 50) and PDFMINER_TEXT_OK:
            try:
                txt = _pdfminer_text(io.BytesIO(context.raw_data)) or ""
                if len(txt) > len(full_text):
                    full_text = _normalise(txt)
                    sections  = [{"title": "Full Document", "text": full_text, "level": 0}]
                    source    = "pdfminer"
            except Exception:
                pass

        # Strategy 3: OCR fallback
        if (not full_text or len(full_text) < 30) and TESSERACT_OK:
            try:
                from pdf2image import convert_from_bytes
                images    = convert_from_bytes(context.raw_data, dpi=150)
                parts     = []
                for i, img in enumerate(images):
                    txt = _tess.image_to_string(img) or ""
                    txt = _normalise(txt)
                    if txt:
                        parts.append(txt)
                        sections.append({"title": f"Page {i+1} (OCR)", "text": txt, "level": 0})
                full_text = "\n\n".join(parts)
                source    = "ocr"
            except Exception:
                pass

        return {"full_text": full_text, "sections": sections,
                "tables": [], "source": source, "page_count": pages}

    # ── DOCX ────────────────────────────────────────────────────────────────
    @staticmethod
    def _read_docx(context: ExtractionContext) -> Dict[str, Any]:
        if not DOCX_OK:
            return {"full_text": "", "sections": [], "tables": [],
                    "source": "docx_unavailable", "page_count": None}
        try:
            doc      = _docx.Document(io.BytesIO(context.raw_data))
            sections : List[Dict] = []
            paras    : List[str]  = []
            tables   : List[List[str]] = []
            current_heading = "Document Start"

            for para in doc.paragraphs:
                txt   = _normalise(para.text)
                style = (para.style.name or "").lower()
                if not txt:
                    continue
                if "heading" in style:
                    level = int(re.search(r"\d", style).group()) if re.search(r"\d", style) else 1
                    current_heading = txt
                    sections.append({"title": txt, "text": "", "level": level})
                else:
                    paras.append(txt)
                    if sections:
                        sections[-1]["text"] += ("\n" if sections[-1]["text"] else "") + txt
                    else:
                        sections.append({"title": current_heading, "text": txt, "level": 0})

            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append(" | ".join(_normalise(c.text) for c in row.cells))
                tables.append(rows)

            full_text = "\n\n".join(
                (s["title"] + "\n" if s["level"] > 0 else "") + s["text"]
                for s in sections
            )
            return {"full_text": full_text, "sections": sections,
                    "tables": tables, "source": "python-docx", "page_count": None}
        except Exception as e:
            return {"full_text": "", "sections": [], "tables": [],
                    "source": f"docx_error: {e}", "page_count": None}

    # ── Spreadsheet ─────────────────────────────────────────────────────────
    @staticmethod
    def _read_spreadsheet(context: ExtractionContext, ext: str) -> Dict[str, Any]:
        tables: List[List[str]]  = []
        sections: List[Dict]     = []
        all_text: List[str]      = []

        if ext == "xlsx" and OPENPYXL_OK:
            try:
                wb = _openpyxl.load_workbook(io.BytesIO(context.raw_data),
                                              read_only=True, data_only=True)
                for sheet_name in wb.sheetnames:
                    ws   = wb[sheet_name]
                    rows = []
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(c) if c is not None else "" for c in row]
                        rows.append(" | ".join(cells))
                        all_text.append(" ".join(cells))
                    tables.append(rows)
                    sections.append({"title": sheet_name,
                                     "text":  "\n".join(rows), "level": 1})
                return {"full_text": "\n".join(all_text), "sections": sections,
                        "tables": tables, "source": "openpyxl", "page_count": None}
            except Exception:
                pass

        if ext == "xls" and XLRD_OK:
            try:
                wb = _xlrd.open_workbook(file_contents=context.raw_data)
                for sheet_name in wb.sheet_names():
                    ws   = wb.sheet_by_name(sheet_name)
                    rows = []
                    for r in range(ws.nrows):
                        cells = [str(ws.cell_value(r, c)) for c in range(ws.ncols)]
                        rows.append(" | ".join(cells))
                        all_text.append(" ".join(cells))
                    tables.append(rows)
                    sections.append({"title": sheet_name,
                                     "text":  "\n".join(rows), "level": 1})
                return {"full_text": "\n".join(all_text), "sections": sections,
                        "tables": tables, "source": "xlrd", "page_count": None}
            except Exception:
                pass

        # CSV / plain text fallback
        return _ContentReader._read_text(context)

    # ── Raw text ─────────────────────────────────────────────────────────────
    @staticmethod
    def _read_text(context: ExtractionContext) -> Dict[str, Any]:
        for enc in ("utf-8", "utf-16", "latin-1", "cp1252"):
            try:
                txt = context.raw_data.decode(enc)
                txt = _normalise(txt)
                paras = _split_paragraphs(txt)
                sections = [{"title": f"Block {i+1}", "text": p, "level": 0}
                            for i, p in enumerate(paras)]
                return {"full_text": txt, "sections": sections,
                        "tables": [], "source": f"text/{enc}", "page_count": None}
            except Exception:
                continue
        return {"full_text": "", "sections": [], "tables": [],
                "source": "decode_failed", "page_count": None}

    # ── Image OCR ────────────────────────────────────────────────────────────
    @staticmethod
    def _read_image(context: ExtractionContext) -> Dict[str, Any]:
        if not TESSERACT_OK:
            return {"full_text": "", "sections": [], "tables": [],
                    "source": "ocr_unavailable", "page_count": None}
        try:
            img      = _PILImage.open(io.BytesIO(context.raw_data))
            txt      = _tess.image_to_string(img) or ""
            txt      = _normalise(txt)
            paras    = _split_paragraphs(txt)
            sections = [{"title": f"OCR Block {i+1}", "text": p, "level": 0}
                        for i, p in enumerate(paras)]
            return {"full_text": txt, "sections": sections,
                    "tables": [], "source": "tesseract_ocr", "page_count": 1}
        except Exception as e:
            return {"full_text": "", "sections": [], "tables": [],
                    "source": f"ocr_error: {e}", "page_count": 1}


# ════════════════════════════════════════════════════════════════════════════
#  Extractor 1 — FullTextExtractor
# ════════════════════════════════════════════════════════════════════════════

class FullTextExtractor(BaseExtractor):
    """
    Extracts the complete verbatim text content of a document, section by
    section, along with summary statistics. The full_text field is included
    in the report so that downstream systems can read the document's content
    directly without re-opening the file.

    Supports: PDF, DOCX, XLSX, XLS, TXT, MD, CSV, images (via OCR).
    """

    name        = "full_text_extractor"
    version     = "1.0"
    category    = CATEGORY_DOCUMENT
    RELIABILITY = RELIABILITY_HIGH

    _LIMITATIONS = [
        "Scanned PDFs without embedded text depend on Tesseract OCR accuracy.",
        "Complex multi-column layouts may merge columns in extracted text.",
        "Password-protected documents cannot be read.",
        "Very large documents (>200 MB) may be slow.",
    ]
    _FALSE_POSITIVES: List[str] = []
    _FALSE_NEGATIVES = [
        "OCR may miss handwritten text or uncommon fonts.",
        "Encrypted/redacted sections will not be extracted.",
    ]

    def _extract(self, context: ExtractionContext) -> Dict[str, Any]:
        content = _ContentReader.read(context)
        full    = content["full_text"]
        secs    = content["sections"]
        tables  = content["tables"]
        source  = content["source"]
        pages   = content["page_count"]

        if not full:
            return {
                "status":  "unavailable",
                "summary": f"No readable text found (source: {source}).",
                "raw_measurements": {"source": source},
                "evidence": {},
                "supports": [],
                "contradicts": [],
            }

        words       = full.split()
        paragraphs  = _split_paragraphs(full)
        sentences   = _split_sentences(full)
        char_count  = len(full)
        word_count  = len(words)
        unique_words = len(set(w.lower().strip(".,;:!?\"'()") for w in words))

        evidence: Dict[str, Any] = {
            # Full content — the core of this extractor
            "full_text":        full,
            "sections":         [
                {
                    "title":    s.get("title", ""),
                    "level":    s.get("level", 0),
                    "text":     s.get("text", ""),
                    "word_count": len(s.get("text", "").split()),
                    "char_count": len(s.get("text", "")),
                }
                for s in secs
            ],
            "tables":           tables,

            # Statistics
            "stats": {
                "source":            source,
                "page_count":        pages,
                "character_count":   char_count,
                "word_count":        word_count,
                "unique_word_count": unique_words,
                "sentence_count":    len(sentences),
                "paragraph_count":   len(paragraphs),
                "section_count":     len(secs),
                "table_count":       len(tables),
                "avg_words_per_sentence": (
                    round(word_count / len(sentences), 2) if sentences else 0
                ),
                "avg_chars_per_word": (
                    round(char_count / word_count, 2) if word_count else 0
                ),
                "vocabulary_richness": (
                    round(unique_words / word_count, 4) if word_count else 0
                ),
            },

            # Content fingerprints (hashes of full text and first 3 sections)
            "content_hashes": {
                "full_text_sha256":      _sha256(full),
                "full_text_normalised":  _sha256(_normalise(full.lower())),
                "first_500_chars":       _sha256(full[:500]) if full else "",
                "last_500_chars":        _sha256(full[-500:]) if full else "",
            },
        }

        supports: List[str] = []
        if word_count < 10:
            supports.append("Very few words extracted — document may be image-only or empty.")
        if source == "ocr":
            supports.append("Text was extracted via OCR — accuracy depends on image quality.")
        if tables:
            supports.append(f"Document contains {len(tables)} table(s)/sheet(s).")
        if pages:
            supports.append(f"Document spans {pages} page(s).")

        return {
            "status":           "ok",
            "summary":          (
                f"Extracted {word_count:,} words, {len(sentences)} sentences, "
                f"{len(secs)} sections from {source} document."
            ),
            "raw_measurements": {
                "word_count": word_count, "char_count": char_count,
                "section_count": len(secs), "source": source,
            },
            "evidence":    evidence,
            "supports":    supports,
            "contradicts": [],
        }


# ════════════════════════════════════════════════════════════════════════════
#  Extractor 2 — ContentBlockHasher
# ════════════════════════════════════════════════════════════════════════════

class ContentBlockHasher(BaseExtractor):
    """
    Generates individual SHA-256 hashes for every paragraph, table row, and
    heading in the document. Each hash is a separate, independently comparable
    evidence item.

    This is the core engine for Authenticity Scoring:
      • Register a document → store N hashes.
      • Verify a suspect document → generate N' hashes.
      • Compare hash-by-hash → compute a match score.

    A typical document yields 50–300+ block hashes.
    Combined with the other extractors, the full Group Hash set
    easily reaches 300–450 unique hashes for scoring.
    """

    name        = "content_block_hasher"
    version     = "1.0"
    category    = CATEGORY_DOCUMENT
    RELIABILITY = RELIABILITY_HIGH

    _LIMITATIONS = [
        "Block boundaries depend on paragraph/line detection quality.",
        "OCR errors may cause block hashes to differ even for identical content.",
        "Reformatted documents (e.g. Word → PDF) will show different block hashes "
        "even if the readable text is identical.",
    ]
    _FALSE_POSITIVES = [
        "Re-exported documents with reflowed text will show partial hash mismatches.",
    ]
    _FALSE_NEGATIVES: List[str] = []

    def _extract(self, context: ExtractionContext) -> Dict[str, Any]:
        content  = _ContentReader.read(context)
        full     = content["full_text"]
        secs     = content["sections"]
        tables   = content["tables"]

        if not full:
            return {
                "status":  "unavailable",
                "summary": "No text content found to hash.",
                "raw_measurements": {}, "evidence": {},
                "supports": [], "contradicts": [],
            }

        blocks: List[Dict[str, Any]] = []

        # ── 1. Paragraph-level hashes ────────────────────────────────────
        for i, para in enumerate(_split_paragraphs(full)):
            if not para:
                continue
            blocks.append({
                "type":          "paragraph",
                "index":         i,
                "preview":       para[:80] + ("…" if len(para) > 80 else ""),
                "word_count":    len(para.split()),
                "sha256":        _sha256(para),
                "sha256_norm":   _sha256(_normalise(para.lower())),
            })

        # ── 2. Section-level hashes ──────────────────────────────────────
        for i, sec in enumerate(secs):
            title = sec.get("title", "")
            text  = sec.get("text", "")
            if text:
                blocks.append({
                    "type":       "section",
                    "index":      i,
                    "title":      title,
                    "level":      sec.get("level", 0),
                    "preview":    text[:80] + ("…" if len(text) > 80 else ""),
                    "word_count": len(text.split()),
                    "sha256":     _sha256(title + "\n" + text),
                    "sha256_norm":_sha256(_normalise((title + " " + text).lower())),
                })
            if title:
                blocks.append({
                    "type":   "heading",
                    "index":  i,
                    "text":   title,
                    "sha256": _sha256(title),
                })

        # ── 3. Table / spreadsheet row hashes ────────────────────────────
        for sheet_i, sheet in enumerate(tables):
            for row_i, row_text in enumerate(sheet):
                if not row_text.strip():
                    continue
                blocks.append({
                    "type":       "table_row",
                    "sheet":      sheet_i,
                    "row":        row_i,
                    "preview":    row_text[:80],
                    "sha256":     _sha256(row_text),
                    "sha256_norm":_sha256(_normalise(row_text.lower())),
                })

        # ── 4. N-gram rolling window hashes (256-char windows, step 128) ─
        #     This catches copy-paste insertions even mid-paragraph.
        window_hashes: List[str] = []
        step = 128
        wsize = 256
        for start in range(0, max(0, len(full) - wsize), step):
            chunk = full[start: start + wsize]
            window_hashes.append(_sha256(chunk))

        # Summary metrics
        total_blocks = len(blocks)
        unique_hashes = len({b["sha256"] for b in blocks})

        return {
            "status":  "ok",
            "summary": (
                f"Generated {total_blocks} content block hashes "
                f"({unique_hashes} unique) + {len(window_hashes)} rolling-window hashes."
            ),
            "raw_measurements": {
                "block_count":         total_blocks,
                "unique_block_hashes": unique_hashes,
                "window_hash_count":   len(window_hashes),
            },
            "evidence": {
                "content_blocks":   blocks,
                "window_hashes":    window_hashes,
                "group_hash_total": total_blocks + len(window_hashes),
            },
            "supports": [
                f"Document yields {total_blocks + len(window_hashes)} individual "
                "evidence hashes for granular authenticity scoring."
            ],
            "contradicts": [],
        }


# ════════════════════════════════════════════════════════════════════════════
#  Extractor 3 — LinguisticProfiler
# ════════════════════════════════════════════════════════════════════════════

class LinguisticProfiler(BaseExtractor):
    """
    Builds a linguistic fingerprint of the document:
      • Per-sentence hashes (100–500 hashes for a typical document)
      • Readability score (Flesch Reading Ease)
      • Vocabulary richness
      • Sentence length distribution
      • Top-50 word frequency map
      • Punctuation / capitalisation pattern hash

    This fingerprint is unique to the author's writing style. If someone
    rewrites even a few sentences, the sentence hashes will change and
    the Authenticity Score will reflect this.
    """

    name        = "linguistic_profiler"
    version     = "1.0"
    category    = CATEGORY_DOCUMENT
    RELIABILITY = RELIABILITY_MEDIUM

    _LIMITATIONS = [
        "Relies on accurate sentence boundary detection.",
        "OCR errors degrade linguistic accuracy significantly.",
        "Very short documents (<100 words) yield unreliable profiles.",
    ]
    _FALSE_POSITIVES = [
        "Translation between languages will invalidate all linguistic hashes.",
    ]
    _FALSE_NEGATIVES = [
        "Professional paraphrasing may preserve linguistic style while changing content.",
    ]

    def _extract(self, context: ExtractionContext) -> Dict[str, Any]:
        content = _ContentReader.read(context)
        full    = content["full_text"]

        if not full or len(full.split()) < 10:
            return {
                "status":  "unavailable",
                "summary": "Insufficient text for linguistic analysis (<10 words).",
                "raw_measurements": {}, "evidence": {},
                "supports": [], "contradicts": [],
            }

        sentences = _split_sentences(full)
        words     = [w.lower().strip(".,;:!?\"'()[]{}") for w in full.split() if w.strip()]
        word_freq : Dict[str, int] = {}
        for w in words:
            word_freq[w] = word_freq.get(w, 0) + 1

        top_words = sorted(word_freq.items(), key=lambda x: -x[1])[:50]

        # Per-sentence hashes
        sentence_hashes = [
            {
                "index":       i,
                "preview":     s[:60] + ("…" if len(s) > 60 else ""),
                "word_count":  len(s.split()),
                "sha256":      _sha256(s),
                "sha256_norm": _sha256(_normalise(s.lower())),
            }
            for i, s in enumerate(sentences) if s.strip()
        ]

        # Sentence length distribution buckets
        lengths = [len(s.split()) for s in sentences]
        dist = {
            "short_1_to_10":     sum(1 for l in lengths if 1 <= l <= 10),
            "medium_11_to_25":   sum(1 for l in lengths if 11 <= l <= 25),
            "long_26_to_50":     sum(1 for l in lengths if 26 <= l <= 50),
            "very_long_over_50": sum(1 for l in lengths if l > 50),
        }

        # Punctuation pattern fingerprint
        punct_pattern = re.sub(r"[a-zA-Z0-9\s]", "", full)
        punct_hash    = _sha256(punct_pattern[:2000])   # first 2000 punct chars

        # Capitalisation pattern hash
        cap_pattern   = "".join("U" if c.isupper() else ("l" if c.islower() else "x") for c in full[:2000])
        cap_hash      = _sha256(cap_pattern)

        readability   = _readability_score(full)

        unique_words  = len(set(words))
        vocab_richness = round(unique_words / len(words), 4) if words else 0

        return {
            "status":  "ok",
            "summary": (
                f"Linguistic profile: {len(sentence_hashes)} sentence hashes, "
                f"vocab richness {vocab_richness:.3f}, "
                f"Flesch {readability.get('flesch_reading_ease', 'N/A')}."
            ),
            "raw_measurements": {
                "sentence_count":   len(sentences),
                "unique_word_count": unique_words,
                "vocab_richness":   vocab_richness,
                "readability":      readability,
            },
            "evidence": {
                "sentence_hashes":          sentence_hashes,
                "sentence_length_dist":     dist,
                "top_50_words":             dict(top_words),
                "punctuation_pattern_hash": punct_hash,
                "capitalisation_hash":      cap_hash,
                "readability":              readability,
                "vocab_richness":           vocab_richness,
                "linguistic_group_total":   len(sentence_hashes) + 2,   # +2 for punct + cap hashes
            },
            "supports": [
                f"Generated {len(sentence_hashes)} per-sentence hashes for "
                "fine-grained linguistic fingerprinting."
            ],
            "contradicts": [],
        }
